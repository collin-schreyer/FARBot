#!/usr/bin/env python3
"""Text extraction from PDF, DOCX, TXT, XLSX files."""

import logging
import os
from pathlib import Path
from typing import Any, Dict, Tuple

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text from PDF, DOCX, TXT, or XLSX. Returns (text, metadata)."""
    path = Path(file_path)
    ext = path.suffix.lower()
    logger.info(f"    📄 Extracting text: {path.name} ({ext}, {path.stat().st_size / 1024:.0f}KB)")
    metadata = {
        "file_name": path.name,
        "file_type": ext,
        "file_size_mb": round(path.stat().st_size / (1024 * 1024), 3),
        "modification_date": "",
    }

    try:
        import time
        metadata["modification_date"] = time.strftime(
            "%Y-%m-%d %H:%M:%S", time.localtime(path.stat().st_mtime)
        )
    except Exception:
        pass

    text = ""
    try:
        if ext == ".pdf":
            text, pdf_meta = _extract_pdf(file_path)
            metadata.update(pdf_meta)
        elif ext == ".docx":
            text = _extract_docx(file_path)
        elif ext == ".txt":
            logger.info(f"    📝 Reading TXT file...")
            text = path.read_text(encoding="utf-8", errors="ignore")
            logger.info(f"    ✅ TXT: {len(text)} chars")
        elif ext in (".xlsx", ".xls", ".xlsm"):
            text = _extract_excel(file_path)
        else:
            logger.info(f"    📝 Reading as plain text ({ext})...")
            text = path.read_text(encoding="utf-8", errors="ignore")
            logger.info(f"    ✅ Plain text: {len(text)} chars")
    except Exception as e:
        logger.error(f"    ❌ Text extraction failed for {path.name}: {e}")
        text = ""

    metadata["char_count"] = len(text)
    metadata["word_count"] = len(text.split())
    return text, metadata


def _ocr_pdf_with_tesseract(file_path: str, max_pages: int = 25, dpi: int = 200) -> str:
    """
    Last-resort OCR for image-only / scanned PDFs.

    Uses PyMuPDF to rasterize pages to PNG (no extra Python deps) and shells
    out to the system `tesseract` binary (installed via apt in the Dockerfile).
    Capped at `max_pages` to keep runtime predictable.

    Returns extracted text, or empty string if OCR is unavailable / fails.
    """
    import os
    import shutil
    import subprocess
    import tempfile

    if shutil.which("tesseract") is None:
        logger.info("    🛈 tesseract binary not present — skipping OCR fallback")
        return ""

    try:
        import fitz  # pymupdf — already a dep of the primary path
    except ImportError:
        logger.warning("    ⚠️ pymupdf unavailable — cannot rasterize for OCR")
        return ""

    parts = []
    with tempfile.TemporaryDirectory(prefix="srt-ocr-") as tmpdir:
        try:
            doc = fitz.open(file_path)
            n = min(len(doc), max_pages)
            logger.info(f"    🖼️ OCR fallback: rasterizing {n}/{len(doc)} pages at {dpi} DPI")
            zoom = dpi / 72.0
            mat = fitz.Matrix(zoom, zoom)
            for i in range(n):
                try:
                    pix = doc[i].get_pixmap(matrix=mat, alpha=False)
                    png_path = os.path.join(tmpdir, f"page_{i:04d}.png")
                    pix.save(png_path)
                    proc = subprocess.run(
                        ["tesseract", png_path, "stdout", "-l", "eng"],
                        capture_output=True, text=True, timeout=120,
                    )
                    if proc.stdout and proc.stdout.strip():
                        parts.append(proc.stdout)
                except subprocess.TimeoutExpired:
                    logger.warning(f"    ⚠️ tesseract timeout on page {i + 1}")
                except Exception as e:
                    logger.warning(f"    ⚠️ OCR failed on page {i + 1}: {e}")
            doc.close()
        except Exception as e:
            logger.warning(f"    ⚠️ OCR fallback failed: {e}")
            return ""

    text = "\n".join(parts)
    logger.info(f"    ✅ OCR fallback recovered {len(text)} chars from {len(parts)} pages")
    return text


def _extract_pdf(file_path: str) -> Tuple[str, Dict]:
    """Extract text from PDF using pymupdf (fast) with pdfplumber fallback,
    and an OCR last-resort for image-only / scanned PDFs."""
    meta = {}
    text = ""

    # Try pymupdf first — much faster and lighter on memory
    try:
        import fitz  # pymupdf
        doc = fitz.open(file_path)
        total = len(doc)
        meta["page_count"] = total
        logger.info(f"    📖 PDF has {total} pages (using pymupdf)")
        pages = []
        for i, page in enumerate(doc):
            t = page.get_text()
            if t and t.strip():
                pages.append(t)
            if (i + 1) % 50 == 0:
                logger.info(f"    📖 Extracted {i + 1}/{total} pages...")
        doc.close()
        text = "\n".join(pages)
        logger.info(f"    ✅ PDF extraction complete: {len(pages)} pages, {len(text)} chars")
    except ImportError:
        logger.info("    pymupdf not available, falling back to pdfplumber")
    except Exception as e:
        logger.warning(f"    ⚠️ pymupdf failed: {e}, falling back to pdfplumber")

    # Fallback to pdfplumber if pymupdf yielded nothing useful or failed.
    if not text or len(text.replace(" ", "")) < 80:
        try:
            import pdfplumber
            pages = []
            MAX_PAGES = 200
            with pdfplumber.open(file_path) as pdf:
                total = len(pdf.pages)
                meta["page_count"] = total
                logger.info(f"    📖 PDF has {total} pages (using pdfplumber)")
                for i, page in enumerate(pdf.pages[:MAX_PAGES]):
                    try:
                        t = page.extract_text()
                        if t:
                            pages.append(t)
                        if (i + 1) % 50 == 0:
                            logger.info(f"    📖 Extracted {i + 1}/{total} pages...")
                    except Exception as e:
                        logger.warning(f"    ⚠️ Page {i + 1} failed: {e}")
                        continue
                if total > MAX_PAGES:
                    logger.warning(f"    ⚠️ Capped at {MAX_PAGES} pages (PDF has {total})")
            plumber_text = "\n".join(pages)
            if len(plumber_text) > len(text):
                text = plumber_text
                logger.info(f"    ✅ pdfplumber extracted: {len(pages)} pages, {len(text)} chars")
        except Exception as e:
            logger.warning(f"    ⚠️ pdfplumber fallback failed: {e}")

    # Last-resort OCR for scanned / image-only PDFs.
    if not text or len(text.replace(" ", "")) < 80:
        ocr_text = _ocr_pdf_with_tesseract(file_path)
        if ocr_text and len(ocr_text) > len(text):
            text = ocr_text
            meta["ocr_applied"] = True

    return text, meta


def _extract_docx(file_path: str) -> str:
    from docx import Document
    logger.info(f"    📝 Extracting DOCX...")
    doc = Document(file_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text += "\n" + " | ".join(cells)
    logger.info(f"    ✅ DOCX: {len(text)} chars, {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    return text


def _extract_excel(file_path: str) -> str:
    from openpyxl import load_workbook
    logger.info(f"    📊 Extracting Excel...")
    wb = load_workbook(file_path, read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        sheet_rows = 0
        for row in ws.iter_rows(values_only=True):
            vals = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if vals:
                parts.append(" | ".join(vals))
                sheet_rows += 1
        logger.info(f"    📊 Sheet '{ws.title}': {sheet_rows} rows")
    wb.close()
    text = "\n".join(parts)
    logger.info(f"    ✅ Excel: {len(text)} chars, {len(parts)} rows total")
    return text
